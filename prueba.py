import docx 
import glob
import xlsxwriter

#construccion de cada una de las funciones

def consejeroPonente(filename):
	doc = docx.Document(filename)
	x = len(doc.paragraphs)
	for i in range(x):
		par = doc.paragraphs[i].text
		z = par.rsplit(': ')
		#Consejero ponente
		if (z[0] == 'Consejero ponente') or (z[0] == 'Consejera ponente'):
			return z[1]
			break

def nombreDemandado(filename):
	doc = docx.Document(filename)
	x = len(doc.paragraphs)
	for i in range(x):
		par = doc.paragraphs[i].text
		z = par.rsplit(': ')
		#funcionario sentenciado
		if z[0] == 'Demandado':
			global demandado
			demandado = z[1]
			return z[1]
			break

def decision(filename):
	doc = docx.Document(filename)
	x = len(doc.paragraphs)
	for i in range(x):
		par = doc.paragraphs[i].text
		try:
			if 'FALLA' == par or 'F A L L A:' == par or 'R E S U E L V E:' == par or 'RESULVE:' == par or'FALLA:' == par:	
				return doc.paragraphs[i+2].text
				break
		except ValueError:
			return ' '

def fechaSentencia(filename):
	doc = docx.Document(filename)
	x = len(doc.paragraphs)		
	for i in range(x): 
		par = doc.paragraphs[i].text
		z = par.rsplit(': ')
		if z[0] == 'Radicación número':
			return doc.paragraphs[i-2].text
			break

def cargoSentenciado(filename):
	doc = docx.Document(filename)
	x = len(doc.paragraphs)	
	for i in range(x):
		par = doc.paragraphs[i].text.lower()
		par = par.rsplit(' ')
		demandado_div = demandado.rsplit(' ')
		if (demandado_div.lower() in s for s in par) and ('representante' in s for s in par):
			return 'Representante'
			break
		elif (demandado_div.lower() in s for s in par) and ('senador' in s for s in par):
			return 'Senador'
			break
		

def motivoPerdida(filename):
	doc = docx.Document(filename)
	x = len(doc.paragraphs)		
	for i in range(x): 
		par = doc.paragraphs[i].text.lower()
		try:
			if ('señala' in par) or ('declarar' in par) or ('declare' in par) or ('perdida de investidura' in par):
				return doc.paragraphs[i].text
				break
		except TypeError:
			return ' '
			break


#crear libro excel
workbook = xlsxwriter.Workbook('tabla.xlsx')
#crear hoja excel
worksheet = workbook.add_worksheet()

#formatos: negrilla
bold = workbook.add_format({'bold': 1})
#formatos: ajustar texto
adjust_txt = workbook.add_format()
adjust_txt.set_text_wrap()
#formatos: texto centrado
adjust_txt.set_align('center')
adjust_txt.set_align('vcenter')

#headers del excel
worksheet.write('A1', 'Archivo', bold)
worksheet.write('B1', 'Consejero ponente', bold)
worksheet.write('C1', 'Fecha sentencia', bold)
worksheet.write('D1', 'Nombre sentenciado', bold)
worksheet.write('E1', 'Cargo', bold)
worksheet.write('F1', 'Resumen caso', bold)
worksheet.write('G1', 'Decision', bold)

#iniciar cuenta desde abajo de los headers
row = 1
col = 0

#escribir datos en el excel
for file in glob.glob('*.docx'):
	worksheet.write_string(row, col, file, adjust_txt)
	worksheet.write_string(row, col+1, consejeroPonente(file), adjust_txt)
	worksheet.write_string(row, col+2, fechaSentencia(file), adjust_txt)
	worksheet.write_string(row, col+3, nombreDemandado(file), adjust_txt)
	worksheet.write_string(row, col+4, cargoSentenciado(file), adjust_txt)
	worksheet.write_string(row, col+5, motivoPerdida(file), adjust_txt)
	worksheet.write_string(row, col+6, decision(file), adjust_txt)
	row += 1

#crear un formato de tabla
worksheet.add_table(f'A1:G{row}', {'header_row': 0})

#cerrar documento
workbook.close()

